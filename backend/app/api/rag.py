"""RAG knowledge management API — upload, list, delete documents."""

import os
import uuid
import datetime
from pathlib import Path
from fastapi import APIRouter, UploadFile, File, Form, Depends, HTTPException
from sqlalchemy.orm import Session
from app.database import get_db
from app.models.models import KnowledgeDocument
from app.services.rag_service import ingest_documents, query_knowledge

router = APIRouter(prefix="/api/rag", tags=["rag"])

UPLOAD_DIR = Path(__file__).resolve().parents[2] / "data" / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


def _extract_text(file_path: Path, suffix: str) -> str:
    """Extract plain text from supported file formats."""
    if suffix in (".md", ".txt"):
        return file_path.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(file_path))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF 解析失败: {e}")

    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(file_path))
            return "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"DOCX 解析失败: {e}")

    raise HTTPException(status_code=400, detail=f"不支持的文件格式: {suffix}")


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    category: str = Form("custom"),
    tags: str = Form(""),
    db: Session = Depends(get_db),
):
    """Upload a document (MD/TXT/PDF/DOCX) into the RAG knowledge base."""
    suffix = Path(file.filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的格式 {suffix}，仅支持: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    # Save uploaded file
    safe_name = f"{uuid.uuid4().hex[:8]}_{file.filename}"
    save_path = UPLOAD_DIR / safe_name
    content = await file.read()
    save_path.write_bytes(content)

    # Extract text
    text = _extract_text(save_path, suffix)
    if not text.strip():
        raise HTTPException(status_code=400, detail="文件内容为空，无法解析")

    # Derive title
    title = Path(file.filename).stem

    # Save to DB
    doc = KnowledgeDocument(
        title=title,
        source=f"uploads/{safe_name}",
        category=category,
        doc_type="user_upload",
        content=text,
        tags=tags or title,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    # Ingest into ChromaDB
    ingest_result = ingest_documents()

    return {
        "code": 0,
        "data": {
            "id": doc.id,
            "title": doc.title,
            "filename": file.filename,
            "size": len(content),
            "chars": len(text),
            "category": category,
            "ingest": ingest_result,
        },
        "message": "文档上传并索引成功",
    }


@router.get("/documents")
def list_documents(db: Session = Depends(get_db)):
    """List all knowledge documents."""
    docs = db.query(KnowledgeDocument).order_by(KnowledgeDocument.created_at.desc()).all()
    return {
        "code": 0,
        "data": [
            {
                "id": d.id,
                "title": d.title,
                "source": d.source,
                "category": d.category,
                "doc_type": d.doc_type,
                "tags": d.tags,
                "chars": len(d.content or ""),
                "created_at": str(d.created_at),
            }
            for d in docs
        ],
        "message": "ok",
    }


@router.delete("/documents/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db)):
    """Delete a knowledge document."""
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    # Delete physical file if it's an upload
    if doc.source and doc.source.startswith("uploads/"):
        file_path = Path(__file__).resolve().parents[2] / "data" / doc.source
        if file_path.exists():
            file_path.unlink()

    db.delete(doc)
    db.commit()

    # Re-ingest to update ChromaDB index
    ingest_documents()

    return {"code": 0, "data": {"id": doc_id}, "message": "文档已删除"}


@router.post("/search")
def search_knowledge(
    query: str = Form(...),
    top_k: int = Form(4),
):
    """Search the RAG knowledge base."""
    results = query_knowledge(query, top_k=top_k)
    return {"code": 0, "data": results, "message": "ok"}
